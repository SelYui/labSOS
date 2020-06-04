# -*- coding: utf-8 -*-
'''
Модули для работы с различными редакторами текстов
'''

import os, docx, pandas
from docx.shared import Pt, Cm
from odf import teletype, text
from odf.opendocument import load

from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT

from lib import module

#создаем файл с таблицей и ИД для лабораторной работы (№ лабы, границы поля:верх, низ, слева, справа)
def create_docx(numb_lab, top, bottom, left, right):
    doc = docx.Document()   #создаем docx
    sections = doc.sections
    #выставляем поля документа
    for section in sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)
    
    if numb_lab == 2:
        #альбомное расположение страниц
        change_orientation(doc)
    
    #определяем стиль и шрифт документа
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    
    return doc

#сохраняем docx файл  
def save_docx(doc,filename):
    doc.save(filename)

#функция создающая таблицу в docx по типу и заполняющая ее для лабораторной №1
def work_docx_from_lab1(index_student, doc, list_student, mas_id, code_mess, fpo_req_id, micro_addr, fpo_exch_id, type_exch):
    #подпись столбцов
    label_column = module.lab_read(1, 4)
    #всего PP_mas_id пунктов, str_mas_id кодов сообщения
    PP_mas_id = len(mas_id)
    str_mas_id =  sum(int(mas_id[i]) for i in range(PP_mas_id))

    if index_student != 0:
        #добавляем разрыв страницы с ненулевого студента
        doc.paragraphs[index_student*2-1].runs[0].add_break(WD_BREAK.PAGE)

    #добавляем параграф со студентом
    doc.add_paragraph(list_student[index_student])
    #создаем таблицу с параметрами
    table = doc.add_table(rows = str_mas_id + 1, cols = len(label_column))
    #стиль таблицы
    table.style = 'Table Grid'
    table_font = 14     #шрифт таблицы
    #автовыравнивание столбцов по содержимому включено
    table.autofit = True
    
    #заполняем строку названий
    for i in range(len(label_column)):
        cell = table.cell(0, i)
        # записываем в ячейку данные
        cell.text = label_column[i]
        # устанавливаем шрифт записываемой ячейки
        table_cell_param(cell, table_font)
        
    #объединение ячеек первого столбца с заполнением
    column = table.columns[0]
    for i in range(len(mas_id)):
        if i == 0:
            sum_i = 1
            column.cells[i+1].merge(column.cells[i + int(mas_id[i])])
            cell = table.cell(sum_i, 0)
        else:
            sum_i = sum_i + int(mas_id[i-1])
            column.cells[sum_i].merge(column.cells[sum_i + int(mas_id[i])-1])
            cell = table.cell(sum_i, 0)
        # записываем в ячейку данные
        cell.text = str(i+1)
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
    
    #заполняем таблицу данными
    for i in range(str_mas_id):
        # Заполняем столбец кодов сообщений
        cell = table.cell(i + 1, 1)
        cell.text = str(code_mess[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - идентификатор запроса ФПО
        cell = table.cell(i + 1, 2)
        cell.text = str(fpo_req_id[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - адрес в микросхеме
        cell = table.cell(i + 1, 3)
        cell.text = str(micro_addr[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        #Заполняем столбец - идентификатор обмена для ФПО
        cell = table.cell(i + 1, 4)
        cell.text = str(fpo_exch_id[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        #Заполняем столбец - тип обмена
        cell = table.cell(i + 1, 5)
        cell.text = str(type_exch[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
    
    doc.add_paragraph(' ')     #пустая строка перед разрывом страницы

#функция создающая таблицу в docx по типу и заполняющая ее для лабораторной №2
def work_docx_from_lab2(index_student, doc, list_student, system_mass, exch_mass, mumb_ou, form_mass, sub_addr, num_dats, fpo_req, id_data, szo_mass):
    #подпись столбцов
    label_column = module.lab_read(2, 1)
    #параметры таблицы
    num_rows = len(form_mass)+1
    num_cols = len(label_column)
    
    if index_student != 0:
        #добавляем разрыв страницы
        doc.paragraphs[index_student*2-1].runs[0].add_break(WD_BREAK.PAGE)
    
    #добавляем параграф со студентом
    doc.add_paragraph(list_student[index_student])
    #создаем таблицу с параметрами
    table = doc.add_table(rows = num_rows, cols = num_cols)
    #стиль таблицы
    table.style = 'Table Grid'
    table_font = 14     #шрифт таблицы
    #автовыравнивание столбцов по содержимому включено
    table.autofit = True
    
    #заполняем строку названий
    for i in range(len(label_column)):
        cell = table.cell(0, i)
        # записываем в ячейку данные
        cell.text = label_column[i]
        # устанавливаем шрифт записываемой ячейки
        table_cell_param(cell, table_font)
    
    #объединение ячеек первого и третьего столбца с заполнением
    column0 = table.columns[0]
    column2 = table.columns[2]
    i = 1
    while i < num_rows:
        column0.cells[i].merge(column0.cells[i + 1])
        column2.cells[i].merge(column2.cells[i + 1])
        cell0 = table.cell(i, 0)
        cell2 = table.cell(i, 2)
        # записываем в ячейку данные
        cell0.text = system_mass[int((i-1)/2)]
        cell2.text = str(mumb_ou[int((i-1)/2)])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell0, table_font)
        table_cell_param(cell2, table_font)
        
        i = i+2
    
    #заполняем таблицу данными
    for i in range(num_rows-1):
        # Заполняем столбец обменов
        cell = table.cell(i + 1, 1)
        cell.text = str(exch_mass[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - Формат
        cell = table.cell(i + 1, 3)
        cell.text = str(form_mass[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - подадрес
        cell = table.cell(i + 1, 4)
        cell.text = str(sub_addr[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - число слов данных
        cell = table.cell(i + 1, 5)
        cell.text = str(num_dats[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - запрос ФПО
        cell = table.cell(i + 1, 6)
        cell.text = str(fpo_req[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - иднтификатор слов данных
        cell = table.cell(i + 1, 7)
        cell.text = str(id_data[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
        
        # Заполняем столбец - СЗО
        cell = table.cell(i + 1, 8)
        cell.text = str(szo_mass[i])
        # устанавливаем шрифт и выравнивание записываемой ячейки
        table_cell_param(cell, table_font)
    
    doc.add_paragraph(' ')     #пустая строка перед разрывом страницы
    
#устанавливаем параметры ячейки
def table_cell_param(cell, font):
    table_paragraph = cell.paragraphs[0]
    table_run = table_paragraph.runs
    table_font = table_run[0].font
    table_font.size= Pt(font)
    #выравнивание текста в ячейке по центу
    table_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
#альбомная ориентация страницы
def change_orientation(doc):
    current_section = doc.sections[-1]
    new_width, new_height = current_section.page_height, current_section.page_width
    new_section = doc.sections[-1] #doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height

#программа для чтения данных из разных файлов
def read_student_list(NameStudent):
    filename, file_extension = os.path.splitext(NameStudent)        #выделяем расширение файла

    temp = []
    if file_extension == '.odt':
        #читаем odt файлы построчно
        text_odt = load(NameStudent)
        allpara = text_odt.getElementsByType(text.P)
        for i in allpara:
            temp.append(teletype.extractText(i))
        student_list = '\n'.join(temp)
        
    elif (file_extension == '.ods'):
        #читаем excel файлы
        temp = pandas.read_excel(NameStudent, engine='odf', header = None)
        #сначала пытаемся считать первый столбец, если не вышло, то нулевой
        try:
            student_list = temp[1].tolist()
        except:
            student_list = temp[0].tolist()
        temp1 = []
        #удаляем несущиствующие строки столбца
        for i in student_list:
            if str(i) != 'nan':
                temp1.append(i)
        #переделываю массив в сплошной список
        student_list = '\n'.join(temp1)
            
    elif (file_extension == '.xls') or (file_extension == '.xlsx'):
        #читаем excel файлы
        temp = pandas.read_excel(NameStudent, header=None)
        #сначала пытаемся считать первый столбец, если не вышло, то нулевой
        try:
            student_list = temp[1].tolist()
        except:
            student_list = temp[0].tolist()
        temp1 = []
        #удаляем несущиствующие строки столбца
        for i in student_list:
            if str(i) != 'nan':
                temp1.append(i)
        #переделываю массив в сплошной список
        student_list = '\n'.join(temp1)
        
    elif(file_extension == '.docx'):
        #читаем docx файл
        doc_file = docx.Document(NameStudent)
        for i in doc_file.paragraphs:
            #student_list.append(i.text)
            temp.append(i.text)
            student_list = '\n'.join(temp)
        
    elif(file_extension == '.txt'):
        #читаем текстовый файл
        file = open(NameStudent, "r", encoding="utf-8")
        student_list = file.read()
        file.close()
    else:
        student_list = False
            
    return student_list

    